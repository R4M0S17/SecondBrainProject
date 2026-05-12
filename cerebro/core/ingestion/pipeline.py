from __future__ import annotations

import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path

from loguru import logger

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".py", ".docx"}


@dataclass
class Document:
    id: str
    content: str
    source_path: str
    chunk_index: int
    file_modified: float
    metadata: dict = field(default_factory=dict)


class IngestionPipeline:
    def __init__(
        self,
        chunk_size: int = 768,
        chunk_overlap: int = 96,
        min_chunk_size: int = 50,
    ) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size

    def ingest(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            logger.warning("Unsupported file format '{}' — skipping {}", ext, path.name)
            return []

        try:
            mtime = os.path.getmtime(file_path)
            size = os.path.getsize(file_path)
        except OSError as e:
            logger.warning("Cannot access file {}: {}", file_path, e)
            return []

        blocks = self._parse(path, ext)
        if not blocks:
            return []

        full_text = "\n".join(blocks)
        chunks = self._chunk(full_text)

        documents = []
        for i, chunk in enumerate(chunks):
            doc_id = hashlib.sha256(chunk.encode()).hexdigest()
            metadata: dict = {"extension": ext, "size": size}
            documents.append(
                Document(
                    id=doc_id,
                    content=chunk,
                    source_path=str(path.resolve()),
                    chunk_index=i,
                    file_modified=mtime,
                    metadata=metadata,
                )
            )
        return documents

    # --- parsers ---

    def _parse(self, path: Path, ext: str) -> list[str]:
        try:
            if ext == ".pdf":
                return self._parse_pdf(path)
            elif ext == ".docx":
                return self._parse_docx(path)
            else:
                return [self._parse_text(path)]
        except Exception as e:
            logger.warning("Failed to parse {}: {}", path.name, e)
            return []

    def _parse_pdf(self, path: Path) -> list[str]:
        import fitz  # PyMuPDF

        blocks: list[str] = []
        with fitz.open(str(path)) as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text().strip()
                if not text:
                    logger.warning(
                        "Page {} of '{}' appears scanned or empty — skipping",
                        page_num + 1,
                        path.name,
                    )
                    continue
                blocks.append(text)
        return blocks

    def _parse_docx(self, path: Path) -> list[str]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        return [p.text for p in doc.paragraphs if p.text.strip()]

    def _parse_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    # --- chunker ---

    def _chunk(self, text: str) -> list[str]:
        words = text.split()
        if not words:
            return []

        step = self.chunk_size - self.chunk_overlap
        chunks: list[str] = []
        start = 0

        while start < len(words):
            chunk_words = words[start : start + self.chunk_size]
            chunk = " ".join(chunk_words)

            if len(chunk_words) >= self.min_chunk_size:
                chunks.append(chunk)
            elif chunks:
                # Merge small trailing chunk into previous rather than dropping content
                chunks[-1] = chunks[-1] + " " + chunk
            # else: single chunk below min_chunk_size — discard per spec

            start += step

        return chunks
